import streamlit as st
import google.generativeai as genai
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import time
import re
import random

# --- 1. CẤU HÌNH TRANG ---
st.set_page_config(
    page_title="HỆ THỐNG RA ĐỀ THI TIỂU HỌC TOÀN DIỆN",
    page_icon="🏫",
    layout="wide"
)

# --- 2. CSS GIAO DIỆN ---
st.markdown("""
<style>
    /* Tab 1 Style */
    .subject-card { padding: 15px; border: 1px solid #ddd; border-radius: 8px; background: #f9f9f9; text-align: center; margin-bottom: 10px; }
    .stTextArea textarea { font-family: 'Times New Roman'; font-size: 16px; }
    .success-box { padding: 10px; background-color: #d4edda; color: #155724; border-radius: 5px; margin-bottom: 10px; }
    
    /* Tab 2 Style */
    .question-box { background-color: #f0f2f6; padding: 15px; border-radius: 10px; border-left: 5px solid #1565C0; margin-bottom: 10px; }
    
    /* Footer */
    .footer {
        position: fixed; left: 0; bottom: 0; width: 100%;
        background-color: #f1f1f1; color: #333;
        text-align: center; padding: 10px; font-size: 14px;
        border-top: 1px solid #ddd; z-index: 100;
    }
    .content-container { padding-bottom: 60px; }
    
    /* Tiêu đề chính */
    .main-header {
        text-align: center; 
        color: #1565C0; 
        font-weight: bold; 
        font-size: 28px; 
        text-transform: uppercase;
        margin-bottom: 20px;
        padding-bottom: 10px;
        border-bottom: 2px solid #eee;
    }
</style>
""", unsafe_allow_html=True)

# --- 3. IMPORT AN TOÀN ---
try:
    import pypdf
except ImportError:
    st.error("⚠️ Thiếu thư viện 'pypdf'. Vui lòng cài đặt: pip install pypdf")

# --- 4. DỮ LIỆU CSDL (GIỮ NGUYÊN BẢN GỐC) ---
SUBJECTS_DB = {
    "Lớp 1": [("Tiếng Việt", "📚"), ("Toán", "🧮")],
    "Lớp 2": [("Tiếng Việt", "📚"), ("Toán", "🧮"), ("Công nghệ", "🔧")],
    "Lớp 3": [("Tiếng Việt", "📚"), ("Toán", "🧮"), ("Tin học", "💻"), ("Công nghệ", "🔧")],
    "Lớp 4": [("Tiếng Việt", "📚"), ("Toán", "🧮"), ("Khoa học", "🔬"), ("Lịch sử & Địa lí", "🌏"), ("Tin học", "💻"), ("Công nghệ", "🔧")],
    "Lớp 5": [("Tiếng Việt", "📚"), ("Toán", "🧮"), ("Khoa học", "🔬"), ("Lịch sử & Địa lí", "🌏"), ("Tin học", "💻"), ("Công nghệ", "🔧")]
}

# DỮ LIỆU CHƯƠNG TRÌNH HỌC (KHÔI PHỤC ĐẦY ĐỦ)
CURRICULUM_DB = {
    "Lớp 1": {
        "Tiếng Việt": {
            "Học kỳ I": [
                {"Chủ đề": "Làm quen với tiếng việt", "Bài học": "Bài 1A: a, b; Bài 1B: c, o; Bài 1C: ô, ơ; Bài 1D: d, đ; Bài 1E: Ôn tập; Bài 2A: e, ê; Bài 2B: h, i; Bài 2C: g, gh; Bài 2D: k, kh; Bài 2E: Ôn tập; Bài 3A: l, m; Bài 3B: n, nh; Bài 3C: ng, ngh; Bài 3D: u, ư; Bài 3E: Ôn tập; Bài 4A: q - qu, gi; Bài 4B: p - ph; Bài 4C: r , s; Bài 4D: t , th; Bài 4E: Ôn tập"},
                {"Chủ đề": "Học chữ ghi vần", "Bài học": "Bài 5A: ch , tr; Bài 5B: x , y; Bài 5C: ua , ưa , ia; Bài 5D: Chữ thường và chữ hoa; Bài 5E: Ôn tập; Bài 6A: â , ai , ay , ây; Bài 6B: oi , ôi , ơi; Bài 6C: ui, ưi; Bài 6D: uôi, ươi; Bài 6E: Ôn tập; Bài 7A: ao, eo; Bài 7B: au, âu; Bài 7C: êu, iu, ưu; Bài 7D: iêu, yêu, ươu; Bài 7E: Ôn tập; Bài 8A: ă, an, ăn, ân; Bài 8B: on, ôn, ơn; Bài 8C: en, ên, un; Bài 8D: in, iên, yên; Bài 8E: uôn, ươn; Bài 9A: ôn tập; Bài 9B: ôn tập; Bài 9C: ôn tập giữa học kì I; Bài 9D: ôn tập giữa học kì I; Bài 9E: ôn tập giữa học kì I; Bài 10A: at, ăt ât; Bài 10B: ot, ôt, ơt; Bài 10C: et, êt, it; Bài 10D: ut, ưt, iêt; Bài 10E: uôt, ươt; Bài 11A: Ôn tập; Bài 11B: am, ăm, âm; Bài 11C: om, ôm, ơm; Bài 11D: em, êm, im; Bài 11E: um, uôm; Bài 12A: ươm, iêm, yêm; Bài 12B: Ôn tập; Bài 12C: ap, ăp, âp; Bài 12D: op, ôp, ơp; Bài 12E: ep, êp, ip; Bài 13A: up, ươp, iêp; Bài 13B: Ôp tập; Bài 13C: ang, ăng, âng; Bài 13D: ong, ông; Bài 13E: ung, ưng; Bài 14A: iêng, uông, ương; Bài 14B: inh, ênh, anh; Bài 14C: Ôn tập; Bài 14D: ac, ăc, âc; Bài 14E: oc, ôc; Bài 15A: uc, ưc; Bài 15B: ich, êch, ach; Bài 15C: iêc, uôc, ươc; Bài 15D: Ôn tập; Bài 15E: oa, oe; Bài 16A: oai, oay; Bài 16B: oan, oăn; Bài 16C: oat, oăt; Bài 16D: oang, oăng, oanh; Bài 16E: oac, oăc, oach; Bài 17A: Ôn tập; Bài 17B: uê, uy, uơ; Bài 17C: uân, uât, uây; Bài 17D: uyên, uyêt, uyt; Bài 17E: Vần ít dùng; Bài 18: ÔT cuối HK I"}
            ],
            "Học kỳ II": [
                {"Chủ đề": "Trường em & Em là búp măng non & Cuộc sống quanh em & Gia đình em", "Bài học": "Bài 19A: Tới trường; Bài 19B: Ở trường thật thú vị; Bài 19 C: Đường đến trường; Bài 19D: Ngôi trường mới; Bài 20A: Bạn bè tuổi thơ; Bài 20B: Bạn thích đồ chơi gì?; Bài 20C: Em nói lời hay; Bài 20D: Giúp bạn vượt khó; Bài 21A: Những âm thanh kì diệu; Bài 21B: Nước có ở đâu?; Bài 21C: Trẻ thơ và trăng; Bài 21D: Những người bạn bé nhỏ; Bài 22A: Con yêu mẹ; Bài 22B: Tập làm đầu bếp; Bài 22C: Em yêu nhà em; Bài 22D: Bố dạy em thế; Bài 23A: Theo bước em đến trường; Bài 23B: Trường đẹp lắm bạn ơi!; Bài 23C: Chuyện ở trường ở lớp; Bài 23D: Đi học thôi, bạn ơi!; Bài 24A: Bạn trong nhà; Bài 24B: Những chuyến đi thú vị; Bài 24C: Niềm vui tuổi thơ; Bài 24D: Những bài học hay; Bài 25A: Những con vật đáng yêu; Bài 25B: Những bông hoa thơm; Bài 25C: Giúp ích cho đời; Bài 25D: Những con vật thông minh; Bài 26A: Con không còn bé nữa; Bài 26B: Bữa cơm gia đình; Bài 26C: Như những người bạn; Bài 26D: Cháu muốn ông bà vui; Bài 28A: Bạn ở trường; Bài 28B: Học cách vui chơi; Bài 28C: Vui chơi ở trường; Bài 28D: Bài học bổ ích; Bài 29A: Nói dối hại thân; Bài 29B: Đi lại an toàn; Bài 29C: Cùng bạn vui chơi; Bài 29D: Điều em ghi nhớ; Bài 30A: Tình yêu thương; Bài 30B: Cuộc sống của các loài cây; Bài 30C: Lời của loài vật; Bài 30D: Điều em muốn biết; Bài 31A: Người thân một nhà; Bài 31B: Nhớ những ngày vui; Bài 31C: Con ngoan của mẹ; Bài 31D: Nhớ lời bố mẹ dặn; Bài 32A: Em lớn lên rồi; Bài 32B: Làm thế nào để khỏe mạnh?; Bài 32C: Đồ chơi tuổi thơ; Bài 32D: Tình bạn; Bài 33A: Những điều giản dị; Bài 33B: Trẻ em là vốn quý; Bài 33C: Những con vật quanh em; Bài 33D: Quanh em có gì thú vị?; Bài 34A: Con xin lỗi; Bài 34B: Biết ơn cha mẹ; Bài 34C: Con yêu của cha mẹ; Bài 34D: Em được yêu thương + Ôn tập"}
            ]
        },
        "Toán": {
            "Học kỳ I": [
                {"Chủ đề": "Các số từ 0 đến 10", "Bài học": "Các số 0, 1,2,3,4,5 (Tr8); Luyện tập (Tr10); Luyện tập (Tr12); Các số 6,7,8,9,10 (Tr14); Luyện tập (Tr16); Luyện tập (Tr18); Nhiều hơn, ít hơn, bằng nhau (Tr20); Luyện tập (Tr22); So sánh số (Lớn hơn, dấu >); So sánh số ( Bé hơn, Dấu < ); So sánh số ( Bằng nhau, Dấu = ); Luyện tập ( Tr30); Mấy và mấy (Tr32); Luyện tập ( Tr36); Luyện tập chung (Tr38); Luyện tập ( Tr40); Luyện tập ( Tr42)"},
                {"Chủ đề": "Làm quen với một số hình học phẳng", "Bài học": "Luyện tập ( Tr44); Hình vuông, hình tròn, hình TG, hình CN; Luyện tập ( Tr48); Thực hành lắp ghép xếp hình; Luyện tập (Tr52); Luyện tập chung (Tr54); Khối lập phương, khối hộp CN; Vị trí định hướng trong không gian; Phải - trái; Luyện tập chung (Tr100)"},
                {"Chủ đề": "Phép cộng, phép trừ trong phạm vi 10", "Bài học": "Phép cộng trong phạm vi 10 (T56); Luyện tập (Tr58); Thêm vào thì bằng mấy?; Số 0 trong phép cộng; Luyện tập (Tr64); Luyện tập (Tr66); Phép trừ trong phạm vi 10 (T68); Tách ra còn lại mấy?; Luyện tập (Tr72); Số 0 trong phép trừ; Luyện tập (Tr76); Luyện tập (Tr78); Bảng cộng, bảng trừ trong phạm vi 10 (Tr80); Bảng trừ (Tr82); Luyện tập (Tr84); Luyện tập chung (Tr86); Luyện tập (Tr88); Luyện tập (Tr90)"}
            ],
            "Học kỳ II": [
                {"Chủ đề": "Số có hai chữ số & Độ dài và đo độ dài", "Bài học": "Bài 21: Số có hai chữ số; Luyện tập (Tr6); Các số tròn chục; Các số đến 99; Bài 22: Ss số có hai chữ số; Bài 23:Bảng các số từ 1- 100; Bài 24: Luyện tập chung; Bài 25: Dài hơn, ngắn hơn + Cao hơn, thấp hơn; Bài 26: Đơn vị đo dộ dài; Bài 27: TH ước lượng và đo dộ dài; Bài 26: Xăng - ti - mét; Bài 28: Luyện tập chung"},
                {"Chủ đề": "Phép cộng và phép trừ (không nhớ) trong pv 100 & Thời gian: Giờ và lịch", "Bài học": "Bài 29: Phép cộng số có hai chữ số với số có một chữ số; Bài 30: Phép cộng số có hai chữ số với số có hai chữ số; Bài 31: Phép trừ số có hai chữ số với số có một chữ số; Bài 32: Phép trừ số có hai chữ số với số có hai chữ số; Bài 33: Luyện tập; Bài 34: Xem giờ đúng trên đồng hồ; Bài 35: Các ngày trong tuần; Bài 36: TH xem lịch và giờ"}
            ]
        }
    },
    "Lớp 2": {
        "Tiếng Việt": {
            "Học kỳ I": [
                {"Chủ đề": "EM LỚN LÊN TỪNG NGÀY", "Bài học": "Bài 1: Tôi là học sinh lớp 2; Bài 2: Ngày hôm qua đâu rồi; Bài 3: Niềm vui của Bi và Bống; Bài 4: Làm việc thật là vui"},
                {"Chủ đề": "ĐI HỌC VUI SAO", "Bài học": "Bài 5: Em có xinh không; Bài 6: Một giờ học; Bài 7: Cây xấu hổ; Bài 8: Cầu thủ dự bị; Bài 9: Cô giáo lớp em; Bài 10: Thời khóa biểu; Bài 11: Cái trống trường em; Bài 12: Danh sách học sinh; Bài 13: Yêu lắm trường ơi!; Bài 14: Em học vẽ; Bài 15: Cuốn sách của em; Bài 16: Khi trang sách mở ra"},
                {"Chủ đề": "NIỀM VUI TUỔI THƠ", "Bài học": "Bài 17: Gọi bạn; Bài 18: Tớ nhớ cậu; Bài 19: Chữ A và những người bạn; Bài 20: Nhím nâu kết bạn; Bài 21: Thả diều; Bài 22: Tớ là lê - gô; Bài 23: Rồng rắn lên mây; Bài 24: Nặn đồ chơi"},
                {"Chủ đề": "Mái ấm gia đình", "Bài học": "Bài 25: Sự tích hoa tỉ muội; Bài 26: Em mang về yêu thương; Bài 27: Mẹ; Bài 28: Trò chơi của bố; Bài 29: Cánh cửa nhớ bà; Bài 30 Thương ông; Bài 31 Ánh sáng của yêu thương; Bài 32 Chơi chong chóng"}
            ],
            "Học kỳ II": [
                {"Chủ đề": "Vẻ đẹp quanh em", "Bài học": "Bài 1: Chuyện bốn mùa; Bài 2: Mùa nước nổi; Bài 3: Họa mi hót; Bài 4: Tết đến rồi; Bài 5: Giọt nước và biển lớn; Bài 6: Mùa vàng; Bài 7: Hạt thóc; Bài 8: Lũy tre"},
                {"Chủ đề": "Hành trình xanh của em", "Bài học": "Bài 9: Vè chim; Bài 10: Khủng long; Bài 11: Sự tích cây thì là; Bài 12: Bờ tre đón khách; Bài 13: Tiếng chổi tre; Bài 14: Cỏ non cười rồi; Bài 15: Những con sao biển; Bài 16 Tạm biệt cánh cam"},
                {"Chủ đề": "Giao tiếp và kết nối", "Bài học": "Bài 17: Những cách chào độc đáo; Bài 18: Thư viện biết đi; Bài 19: Cảm ơn anh hà mã; Bài 20: Từ chú bồ câu đến in- tơ- nét"},
                {"Chủ đề": "Con người Việt Nam", "Bài học": "Bài 21: Mai An Tiêm; Bài 22: Thư gửi bố ngoài đảo; Bài 23: Bóp nát quả cam; Bài 24: Chiếc rễ đa tròn"},
                {"Chủ đề": "Việt Nam quê hương em", "Bài học": "Bài 25: Đất nước chúng mình; Bài 26: Trên các miền đất nước; Bài 27: Chuyện quả bầu; Bài 28: Khám phá đáy biển ở Trường Sa; Bài 29 Hồ Gươm; Bài 30: Cánh đồng quê em"}
            ]
        },
        "Toán": {
            "Học kỳ I": [
                {"Chủ đề": "Ôn tập và bổ sung", "Bài học": "Bài 1: Ôn tập các số đến 100; Bài 2: Tia số. Số liền trước, số liền sau; Bài 3: Số hạng. Tổng/Số bị trừ, số trừ, hiệu; Bài 4: Hơn, kém nhau bao nhiêu; Bài 5: Luyện tập; Bài 6: Luyện tập chung"},
                {"Chủ đề": "Phép cộng, phép trừ trong phạm vi 20", "Bài học": "Bài 7: Phép cộng (qua 10) trong pv 20; Bài 8: Bảng cộng (qua 10); Bài 9: Giải bài toán về thêm một số đv/bớt một số đv; Bài 10: Luyện tập; Bài 11: Phép trừ (qua 10) trong pv 20; Bài 12: Bảng trừ (qua 10); Bài 13: Giải BT về nhiều hơn/ít hơn một số đv; Bài 14: Luyện tập chung"},
                {"Chủ đề": "Làm quen với khối lượng, dung tích", "Bài học": "Bài 15: Nặng hơn, nhẹ hơn/Ki - lô- gam; Bài 16: Lít; Bài 17: Thực hành và trải nghiệm với các đơn vị Ki - lô - gam, Lít; Bài 18: Luyện tập chung"},
                {"Chủ đề": "Phép cộng, phép trừ có nhớ trong phạm vi 100", "Bài học": "Bài 19: Phép cộng (có nhớ) số có hai chữ số với số có một chữ số; Bài 20: Phép cộng (có nhớ) số có hai chữ số với số có hai chữ số; Bài 21: Luyện tập; Bài 22: Phép trừ (có nhớ) số có hai chữ số với số có một chữ số; Bài 23: Phép trừ (có nhớ) số có 2 chữ số với số có hai chữ số; Bài 24: Luyện tập chung"},
                {"Chủ đề": "Làm quen với hình phẳng", "Bài học": "Bài 25: Điểm, đoạn thẳng/Đường thẳng, đường cong, ba điểm thẳng hàng; Bài 26: Đường gấp khúc/Hình tứ giác; Bài 27: Thực hành gấp, cắt, ghép, xếp hình/Vẽ đoạn thẳng; Bài 28: Luyện tập chung"},
                {"Chủ đề": "Ngày - giờ, giờ - phút, Xem đồng hồ, xem lịch ngày- tháng.", "Bài học": "Bài 29: Ngày - giờ, giờ - phút/Xem đồng hồ; Bài 30: Ngày – tháng; Bài 31: Thực hành và trải nghiệm xem đồng hồ, xem lịch; Bài 32: Luyện tập chung"},
                {"Chủ đề": "Ôn tập học kì I", "Bài học": "Bài 33: ÔT phép cộng, phép trừ trong pv 20; Bài 33: Ôn tập phép cộng, phép trừ trong phạm vi 100; Bài 34: Luyện tập; Bài 35: Luyện tập; Bài 36: Luyện tập"}
            ],
            "Học kỳ II": [
                {"Chủ đề": "Phép nhân, phép chia", "Bài học": "Bài 37: Phép nhân; Bài 38: Thừa số, tích; Bài 39: Bảng nhân 2; Bài 40: Bảng nhân 5; Bài 41: Phép chia; Bài 42: Số bị chia, cố chia, thương; Bài 43: Bảng chia 2; Bài 44: Bảng chia 5; Bài 45: Luyện tập chung"},
                {"Chủ đề": "Làm quen với hình khối", "Bài học": "Bài 46: Khối trụ, khối cầu; Bài 47: Luyện tập chung"},
                {"Chủ đề": "Các số trong phạm vi 1000", "Bài học": "Bài 48: Đơn vị, chục, trăm, nghìn; Bài 49: Các số tròn trăm/Các số tròn chục; Bài 50: So sánh các số tròn trăm, tròn chục; Bài 51: Số có ba chữ số; Bài 52: Viết số thành tổng các trăm, chục, đơn vị; Bài 53: So sánh các số có ba chữ số; Bài 54: Luyện tập chung"},
                {"Chủ đề": "Độ dài và đơn vị đo độ dài. Tiền VN", "Bài học": "Bài 55: Đề - xi - mét/Mét/Ki-lô-mét; Bài 56: Giới thiệu Tiền Việt Nam; Bài 57: Thực hành và trải nghiệm đo độ dài; Bài 58: Luyện tập chung"},
                {"Chủ đề": "Phép cộng, phép trừ trong phạm vi 1000", "Bài học": "Bài 59: Phép cộng (không nhớ); Bài 60: Phép cộng (có nhớ); Bài 61: Phép trừ (không nhớ); Bài 62: Phép trừ (có nhớ); Bài 63: Luyện tập chung"},
                {"Chủ đề": "Làm quen với thống kê xuất sắc", "Bài học": "Bài 64: Thu thập, phân loại, kiểm đếm số liệu; Bài 65: Biểu đồ tranh; Bài 66: Chắc chắn, có thể, không thể; Bài 67: Thực hành và trải nghiệm thu thập, phân loại, kiểm đếm số liệu"},
                {"Chủ đề": "Ôn tập cuối năm", "Bài học": "Bài 68: Ôn tập các số trong phạm vi 1000; Bài 69: Ôn tập phép cộng. phép trừ trong phạm vi 100; Bài 70: ÔT phép +, phép - trong pv 1000; Bài 71: ÔT phép nhân, phép chia; Bài 72: Ôn tập hình học; Bài 73: ÔT đo lường; Bài 74: ÔTKT số liệu và lựa chọn KN; Bài 75: Ôn tập chung"}
            ]
        },
        "Công nghệ": {
            "Học kỳ I": [
                {"Chủ đề": "Công nghệ và đời sống", "Bài học": "Bài 1: Lợi ích của hoa, cây cảnh đối với đời sống; Bài 2: Một số loại hoa, cây cảnh phổ biến; Bài 3: Vật liệu và dụng cụ trồng hoa, cây; Bài 4: Gieo hạt hoa, cây cảnh trong chậu; Bài 5: Trồng hoa, cây cảnh trong chậu; Bài 6: Chăm sóc hoa, cây cảnh trong chậu"}
            ],
            "Học kỳ II": [
                {"Chủ đề": "Thủ công kĩ thuật", "Bài học": "Bài 7: Giới thiệu bộ lắp ghép mô hình kĩ thuật; Bài 8: Lắp ghép mô hình bập bênh; Bài 9: Lắp ghép mô hình robot; Bài 10: Đồ chơi dân gian; Bài 11: Làm đèn lồng; Bài 12: Làm chuồn chuồn thăng bằng"}
            ]
        }
    },
    "Lớp 3": {
        "Tiếng Việt": {
            "Học kỳ I": [
                {"Chủ đề": "Những trải nghiệm thú vị", "Bài học": "B1: Ngày gặp lại; B2: Về thăm quê; B3: Cánh rừng trong nắng; B4: Lần đầu ra biển; B5: Nhật kí tập bơi; B6: Tập nấu ăn; B7: Mùa hè lấp lánh; Bài 8: Tạm biệt mùa hè"},
                {"Chủ đề": "Công trường rộng mở", "Bài học": "B9: Đi học vui sao; Bài 10: Con đường tới trường; Bài 11: Lời giải toán đặc biệt; Bài 12: Bài tập làm văn; Bài 13: Bàn tay cô giáo; Bài 14: Cuộc họp của chữ viết; Bài 15: Thư viện; Bài 16: Ngày em vào đội"},
                {"Chủ đề": "Mái nhà yêu thương", "Bài học": "B17: Ngưỡng cửa; Bài 18: Món quà đặc biệt; Bài 19: Khi cả nhà bé tí; Bài 20: Trò chuyện cùng mẹ; Bài 21: Tia nắng bé nhỏ; Bài 22: Để cháu nắm tay ông; Bài 23: Tôi yêu em tôi; Bài 24: Bạn nhỏ trong nhà"},
                {"Chủ đề": "Mái ấm gia đình", "Bài học": "Bài 25: Những bậc đá chạm mây; Bài 26: Đi tìm mặt trời; B27: Những chiếc áo ấm; Bài 28: Con đường của bé; Bài 29: Ngôi nhà trong cỏ; Bài 30: Những ngọn hải đăng; Bài 31: Người làm đồ chơi; Bài 32: Cây bút thần"}
            ],
            "Học kỳ II": [
                {"Chủ đề": "Những sắc màu TN", "Bài học": "Bài 1: Bầu trời; Bài 2: Mưa; Bài 3: Cóc kiện trời; Bài 4: Những cái tên đáng yêu; Bài 5: ngày hội rừng xanh; Bài 6: Cây gạo; Bài 7: Mặt trời xanh của tôi; Bài 8: Bầy voi rừng trường sơn"},
                {"Chủ đề": "Bài học từ cuộc sống", "Bài học": "Bài 9: Lời kêu gọi toàn dân tập thể dục; Bài 10: Quả hồng của thỏ con; Bài 11: Chuyện bên cửa sổ; Bài 12: Tay trái và tay phải; Bài 13: Mèo đi câu cá; Bài 14: Học nghề; Bài 15: Ngày như thế nào là đẹp?; Bài 16: A lô, tớ đây"},
                {"Chủ đề": "Đất nước ngàn năm", "Bài học": "Bài 17: Đất nước là gì?; Bài 18: Núi quê tôi; Bài 19: Sông hương; Bài 20: Tiếng nước mình; Bài 21: Nhà rông; Bài 22: Sự tích ông đùng, bà đùng; Bài 23: Hai bà trưng; Bài 24: Cùng bác qua suối"},
                {"Chủ đề": "Trái đất của chúng mình", "Bài học": "Bài 25: Ngọn lửa ô - lim - pích; Bài 26: Rô - bốt ở quanh ta; Bài 27: Thư của ông trái đất gửi các bạn nhỏ; B28: Những điều nhỏ tớ làm cho trái đất; Bài 29: Bác sĩ y- éc- xanh; Bài 30: Một mái nhà chung"}
            ]
        },
        "Toán": {
            "Học kỳ I": [
                {"Chủ đề": "Ôn tập và bổ sung", "Bài học": "Bài 1: Luyện tập (trang 6); Bài 2: Luyện tập (trang 9); Bài 3: Tìm số hạng trong một tổng/Tìm số bị trừ, số trừ; Bài 4: Ôn tập bảng nhân 2, 5/bảng chia 2, 5; Bài 5: Bảng nhân 3/Bảng chia 3"},
                {"Chủ đề": "Bảng nhân, bảng chia", "Bài học": "Bài 6: Bảng nhân 4/Bảng chia 4; Bài 9: Bảng nhân 6/bảng chia 6; Bài 10: Bảng nhân 7/bảng chia 7; Bài 11: Bảng nhân 8/bảng chia 8; Bài 12: Bảng nhân 9/bảng chia 9; Bài 13: Tìm thừa số trong một tích/Tìm số bị chia, số chia; Bài 14: Một phần mấy; Bài 15: Luyện tập"},
                {"Chủ đề": "Làm quen với hình phẳng, hình khối", "Bài học": "Bài 16: Điểm ở giữa, trung điểm của đoạn thẳng; Bài 17: Hình tròn. Tâm, bán kính, đường kính; Bài 18: Góc, góc vuông, góc không vuông; Bài 19: Hình tam giác, hình tứ giác/Hình chữ nhật, hình vuông; Bài 20: Thực hành vẽ góc vuông, vẽ đường tròn, hình vuông, hình chữ nhật; Bài 21: Khối lập phương, khối hộp chữ nhật; Bài 22: Luyện tập"},
                {"Chủ đề": "Phép nhân, phép chia trong phạm vi 100", "Bài học": "Bài 23: Nhân số có 2 chữ số với số có một chữ số; Bài 24: Gấp một số lên một số lần; Bài 25: Phép chia hết, phép chia có dư; Bài 26: Chia số có 2 chữ số cho số có một chữ số; Bài 27: Giảm một số đi một số lần; Bài 28: Bài toán giải bằng hai phép tính; Bài 29: Luyện tập"}
            ],
            "Học kỳ II": [
                {"Chủ đề": "Một số đơn vị đo độ dài, khối lượng, dung tích, nhiệt độ", "Bài học": "Bài 30: Mi – li – mét; Bài 31: Gam; Bài 32: Mi - li - lít; Bài 33: Nhiệt độ/Đv đo nhiệt độ; Bài 34: Thực hành và trải nghiệm; Bài 35: Luyện tập"},
                {"Chủ đề": "Phép nhân, phép chia trong phạm vi 1000", "Bài học": "Bài 36: Nhân số có ba chữ số với số có một chữ số; Bài 37: Chia số có 3 chữ số cho số có một chữ số; Bài 38: Làm quen với biểu thức/Tính giá trị của biểu thức; Bài 39: So sánh số lớn gấp mấy lần số bé; Bài 40: Luyện tập"},
                {"Chủ đề": "Ôn tập và Ôn tập cuối năm", "Bài học": "Bài 45: Số có 4 chữ số/Số 10.000; Bài 47: Làm quen với chữ số La Mã; Bài 48: Làm tròn số đến hàng chục, hàng trăm; Bài 50: Chu vi hình tam giác, hình tứ giác; Bài 51: Diện tích của một hình/Xăng ti mét vuông; Bài 52: DT hình chữ nhật/DT hình vuông; Bài 54: Phép cộng trong phạm vi 10 000; Bài 55: Phép trừ trong phạm vi 10 000; Bài 56: Nhân số có 4 chữ số cho số có một chữ số; Bài 57: Chia số có 4 chữ số cho số có một chữ số; Bài 59: Số có 5 chữ số/Số 100 000; Bài 60: So sánh các số trong pv 100 000; Bài 61: Làm tròn các số đến hàng nghìn, hàng chục nghìn; Bài 63: Phép cộng trong phạm vi 100 000; Bài 64: Phép trừ trong phạm vi 100 000; Bài 66: Xem đồng hồ. Tháng – năm; Bài 68: Tiền Việt Nam; Bài 70: Nhân số có 5 chữ số với số có một chữ số; Bài 71: Chia số có năm chữ số cho số có một chữ số; Bài 73: Thu thập, phân loại, ghi chép số liệu. bảng số liệu; Bài 74: Khả năng xảy ra của một sự kiện"}
            ]
        },
        "Tin học": {
            "Học kỳ I": [
                {"Chủ đề": "Máy tính và em", "Bài học": "Bài 1. Thông tin và quyết định; Bài 2. Xử lí thông tin; Bài 3. Máy tính và em; Bài 4. Làm việc với máy tính; Bài 5. Sử dụng bàn phím"},
                {"Chủ đề": "Mạng máy tính và Internet", "Bài học": "Bài 6. Khám phá thông tin trên Internet"},
                {"Chủ đề": "Tổ chức lưu trữ, tìm kiếm và trao đổi thông tin", "Bài học": "Bài 7. Sắp xếp để dễ tìm; Bài 8. Sơ đồ hình cây. Tổ chức thông tin trong máy tính; Bài 9. Thực hành với tệp và thư mục trong máy tính"},
                {"Chủ đề": "Đạo đức, pháp luật và văn hoá trong môi trường số", "Bài học": "Bài 10. Bảo vệ thông tin khi dùng máy tính"}
            ],
            "Học kỳ II": [
                {"Chủ đề": "Ứng dụng tin học", "Bài học": "Bài 12. Tìm hiểu về thế giới tự nhiên; Bài 11. Bài trình chiếu của em"},
                {"Chủ đề": "Giải quyết vấn đề với sự trợ giúp của máy tính", "Bài học": "Bài 14. Em thực hiện công việc như thế nào?; Bài 15. Công việc được thực hiện theo điều kiện; Bài 16. Công việc của em và sự trợ giúp của máy tính"}
            ]
        },
        "Công nghệ": {
            "Học kỳ I": [
                {"Chủ đề": "Công nghệ và đời sống", "Bài học": "Bài 1: Tự nhiên và công nghệ (HĐ1, HĐ2, HĐ3); Bài 2: Sử dụng đèn học (HĐ 1, HĐ 2, HĐ 3); Bài 3: Sử dụng quạt điện (HĐ1, HĐ2, HĐ3); Bài 4: Sử dụng máy thu thanh (HĐ1, HĐ2, HĐ3, HĐ4); Bài 5: Sử dụng máy thu hình (HĐ1, HĐ2, HĐ3, HĐ4); Bài 6: An toàn với môi trường công nghệ trong gia đình (HĐ1, HĐ2, HĐ3)"}
            ],
            "Học kỳ II": [
                {"Chủ đề": "Thủ công kĩ thuật", "Bài học": "Bài 7: Dụng cụ và vật liệu làm thủ công; Bài 8: Làm đồ dùng học tập; Bài 9: Làm biển báo giao thông; Bài 10: Làm đồ chơi; Bài 11: Làm đèn lồng; Bài 12: Làm chuồn chuồn thăng bằng"}
            ]
        }
    },
    "Lớp 4": {
        "Tiếng Việt": {
            "Học kỳ I": [
                {"Chủ đề": "Mỗi người một vẻ", "Bài học": "Bài 1: Điều kì diệu; Bài 2: Thi nhạc; Bài 3: Anh em sinh đôi; Bài 4: Công chúa và người dẫn chuyện; Bài 5: Thằn lằn xanh và tắc kè; Bài 6: Nghệ sĩ trống; Bài 7: Những bức chân dung; Bài 8: Đò ngang"},
                {"Chủ đề": "Trải nghiệm và khám phá", "Bài học": "Bài 9: Bầu trời trong quả trứng; Bài 10: Tiếng nói của cỏ cây; Bài 11: Tập làm văn; Bài 12: Nhà phát minh 6 tuổi; Bài 13: Con vẹt xanh; Bài 14: Chân trời cuối phố; Bài 15: Gặt chữ trên non; Bài 16: Trước ngày xa quê"},
                {"Chủ đề": "Niềm vui sáng tạo", "Bài học": "Bài 17: Vẽ màu; Bài 18: Đồng cỏ nở hoa; Bài 19: Thanh âm của núi; Bài 20: Bầu trời mùa thu; Bài 21: Làm thỏ con bằng giấy; Bài 22: Bức tường có nhiều phép lạ"},
                {"Chủ đề": "Chắp cánh ước mơ", "Bài học": "Bài 23: Bét -tô - ven và Bản xô –nát ánh trăng; Bài 24: Người tìm đường lên các vì sao; Bài 25: Bay cùng ước mơ; Bài 26: Con trai người làm vườn; Bài 27: Nếu em có một khu vườn; Bài 28: Bốn mùa mơ ước; Bài 29: Ở vương quốc tương lai; Bài 30: Cánh chim nhỏ; Bài 31: Nếu chúng mình có phép lạ; Bài 32: Anh Ba"}
            ],
            "Học kỳ II": [
                {"Chủ đề": "Sống để yêu thương", "Bài học": "Bài 1: Hải thượng Lãn Ông; Bài 2: Vệt phấn trên mặt bàn; Bài 3: Ông bụt đã đến; Bài 4: Quả ngọt cuối mùa; Bài 5: Tờ báo tường của tôi; Bài 6: Tiếng ru; Bài 7: Con muốn làm một cái cây; Bài 8: Trên khám tre đầu ngõ"},
                {"Chủ đề": "Uống nước nhớ nguồn", "Bài học": "Bài 9: Sự tích con Rồng, cháu Tiên; Bài 10: Cảm xúc Trường Sa; Bài 11: Sáng tháng Năm; Bài 12: Chàng trai Làng Phù Ủng; Bài 13: Vườn của ông tôi; Bài 14: Trong lời mẹ hát; Bài 15: Người thầy đầu tiên của bố tôi; Bài 16: Ngựa biên phòng"},
                {"Chủ đề": "Quê hương trong tôi", "Bài học": "Bài 17: Cây đa quê hương; Bài 18: Bước mùa xuân; Bài 19: Đi hội Chùa Hương; Bài 20: Chiều ngoại ô; Bài 21: Những cánh buồm; Bài 22: Cái câu; Bài 23: Đường đi Sa Pa; Bài 24: Quê ngoại"},
                {"Chủ đề": "Vì một thế giới bình yên", "Bài học": "Bài 25: Khu bảo tồn động vật hoang dã Ngô rông- gô – rô; Bài 26: Ngôi nhà của yêu thương; Bài 27: Băng tan; Bài 28: Chuyến du lịch thú vị; Bài 29: Lễ hội ở Nhật Bản; Bài 30: Ngày hội"}
            ]
        },
        "Toán": {
            "Học kỳ I": [
                {"Chủ đề": "Số có nhiều chữ số", "Bài học": "Bài 10: Số có sáu chữ số. Số 1000000; Bài 11: Hàng và lớp; Bài 12: Các số trong phạm vi lớp triệu; Bài 13: Làm tròn số đến hàng trăm nghìn; Bài 14: So sánh các số có nhiều chữ số; Bài 15: Làm quen với dãy số tự nhiên; Bài 16: Luyện tập chung"},
                {"Chủ đề": "Một số đơn vị đo đại lượng", "Bài học": "Bài 17: Yến, tạ, tấn; Bài 18: Đề- xi- mét vuông, mét vuông, Mi- li- mét vuông; Bài 19: Giây, thế kỉ; Bài 20: Thực hành và trải nghiệm sử dụng một số đơn vị đo đại lượng; Bài 21: Luyện tập chung"},
                {"Chủ đề": "Phép cộng và phép trừ", "Bài học": "Bài 22: Phép cộng các số có nhiều chữ số; Bài 23: Phép trừ các số có nhiều chữ số; Bài 24: Tính chất giao hoán và kết hợp của phép cộng; Bài 25: Tìm hai số khi biết tổng và hiệu của hai số đó; Bài 26: Luyện tập chung"}
            ],
            "Học kỳ II": [
                {"Chủ đề": "Phép nhân, phép chia", "Bài học": "Bài 38: Nhân với số có một chữ số; Bài 39: Chia cho số có một chữ số; Bài 40: Tính chất giao hoán và kết hợp của phép nhân; Bài 41: Nhân, chia với 10,100,1 000; Bài 42: Tính chất phân phối của phép nhân đối với phép cộng; Bài 43: Nhân với số có hai chữ số; Bài 44: Chia cho số có hai chữ số; Bài 45: Thực hành và trải nghiệm ước lượng trong tính toán; Bài 46: Tìm số trung bình cộng; Bài 47: Bài toán liên quan đến rút về đơn vị; Bài 48: Luyện tập chung"},
                {"Chủ đề": "Phân số, khái niệm phân số", "Bài học": "Bài 53: Khái niệm phân số; Bài 54: Phân số và phép chia số tự nhiên; Bài 55: Tính chất cơ bản của phân số; Bài 56: Rút gọn phân số; Bài 57: Quy đồng mẫu số các phân số; Bài 58: So sánh phân số; Bài 59: Luyện tập chung"},
                {"Chủ đề": "Phép cộng, phép trừ phân số", "Bài học": "Bài 60: Phép cộng phân số; Bài 61: Phép trừ phân số; Bài 62: Luyện tập chung"},
                {"Chủ đề": "Phép nhân, phép chia phân số", "Bài học": "Bài 63: Phép nhân phân số; Bài 64: Phép chia phân số; Bài 65: Tìm phân số của một số; Bài 66: Luyện tập chung"},
                {"Chủ đề": "Ôn tập cuối năm", "Bài học": "Bài 67: Ôn tập số tự nhiên; Bài 68: Ôn tập phép tính với số tự nhiên; Bài 69: Ôn tập phân số; Bài 70: Ôn tập phép tính với phân số; Bài 71: Ôn tập hình học và đo lường; Bài 72: Ôn tập một số yếu tố thống kê và xác suất; Bài 73: Ôn tập chung"}
            ]
        },
        "Lịch sử & Địa lí": {
            "Học kỳ I": [
                {"Chủ đề": "ĐỊA PHƯƠNG EM", "Bài học": "Bài 2. Thiên nhiên và con người ở địa phương em; Bài 3. Lịch sử và văn hoá truyền thống địa phương em"},
                {"Chủ đề": "TRUNG DU VÀ VÙNG NÚI BẮC BỘ", "Bài học": "Bài 4: Thiên nhiên vùng Trung du và miền núi Bắc bộ; Bài 5: Dân cư, hoạt động sản xuất ở vùng Trung du và miền núi Bắc bộ; Bài 6: Một số nét văn hóa ở vùng Trung du và miền núi Bắc bộ; Bài 7: Đền Hùng và lễ giỗ Tổ Hùng Vương"},
                {"Chủ đề": "ĐỒNG BẰNG BẮC BỘ", "Bài học": "Bài 8: Thiên nhiên vùng đồng bằng Bắc Bộ; Bài 9: Dân cư, hoạt động sản xuất ở vùng Đồng bằng Bắc Bộ; Bài 10: Một số nét văn hóa ở vùng Đồng bằng Bắc Bộ; Bài 11: Sông Hồng và văn minh sông Hồng; Bài 12: Thăng Long – Hà Nội; Bài 13: Văn Miếu – Quốc tử giám"}
            ],
            "Học kỳ II": [
                {"Chủ đề": "DUYÊN HẢI MIỀN TRUNG", "Bài học": "Bài 15: Thiên nhiên vùng duyên hải miền Trung; Bài 16: Dân cư, hoạt động sản xuất ở vùng duyên hải miền Trung; Bài 17: Một số nét văn hóa ở vùng duyên hải miền Trung; Bài 18: Cố đô Huế; Bài 19: Phố cổ Hội An"},
                {"Chủ đề": "TÂY NGUYÊN", "Bài học": "Bài 20: Thiên nhiên vùng Tây Nguyên; Bài 21: Dân cư, hoạt động sản xuất ở vùng Tây Nguyên; Bài 22: Một số nét văn hóa và truyền thống yêu nước, cách mạng của đồng bào Tây Nguyên; Bài 23: Lễ hội cồng chiêng Tây Nguyên"},
                {"Chủ đề": "NAM BỘ", "Bài học": "Bài 24: Thiên nhiên vùng Nam Bộ; Bài 25: Dân cư, hoạt động sản xuất vùng Nam Bộ; Bài 26: Một số nét văn hóa và truyền thống yêu nước, cách mạng của đồng bào Nam Bộ; Bài 27: Thành phố Hồ Chí Minh; Bài 28: Địa đạo củ chi"}
            ]
        },
        "Khoa học": {
            "Học kỳ I": [
                {"Chủ đề": "CHẤT", "Bài học": "Bài 1: Thành phần và vai trò của đất đối với cây trồng; Bài 2: Ô nhiễm, xói mòn đất và bảo vệ môi trường đất; Bài 3: Hỗn hợp và dung dịch; Bài 4: Đặc điểm của chất ở trạng thái rắn, lỏng, khí. Sự biến đổi trạng thái của chất; Bài 5: Sự biến đổi hóa học của chất; Bài 6: Ôn tập chủ đề chất"},
                {"Chủ đề": "NĂNG LƯỢNG", "Bài học": "Bài 7: Vai trò của năng lượng; Bài 8: Sử dụng năng lượng điện; Bài 9: Mạch điện đơn giản, vật dẫn điện và vật cách điện; Bài 10: Năng lượng chất đốt; Bài 11: Sử dụng năng lượng mặt trời, năng lượng gió, năng lượng nước chảy; Bài 12: Ôn tập chủ đề năng lượng"},
                {"Chủ đề": "THỰC VẬT VÀ ĐỘNG VẬT", "Bài học": "Bài 13: Sinh sản của thực vật có hoa; Bài 14: Sự phát triển của cây con; Bài 15: Sinh sản của thực vật có hoa; Bài 16: Vòng đời và sự phát triển của động vật; Bài 17: ôn tập chủ đề thực vật và động vật"}
            ],
            "Học kỳ II": [
                {"Chủ đề": "VI KHUẨN", "Bài học": "Bài 18: Vi khuẩn xung quanh chúng ta; Bài 19: Vi khuẩn có ích trong chế biến thực phẩm; Bài 20: Vi khuẩn gây bệnh ở người và cách phòng tránh; Bài 21: Ôn tập chủ đề vi khuẩn"},
                {"Chủ đề": "CON NGƯỜI VÀ SỨC KHỎE", "Bài học": "Bài 22: Sự hình thành cơ thể người; Bài 23: Các giai đoạn phát triển chính của con người; Bài 24: Nam và nữ; Bài 25: Chăm sóc sức khoẻ tuổi dậy thì; Bài 26: Phòng tránh bị xâm hại; Bài 27: Ôn tập chủ đề con người và sức khoẻ"},
                {"Chủ đề": "SINH VẬT VÀ MÔI TRƯỜNG", "Bài học": "Bài 28: Chức năng của môi trường đối với sinh vật; Bài 29: Tác động của con người và một số biện pháp bảo vệ môi trường; Bài 30: ôn tập chủ đề sinh vật và môi trường"}
            ]
        },
        "Tin học": {
            "Học kỳ I": [
                {"Chủ đề": "MÁY TÍNH VÀ EM", "Bài học": "Bài 1. Em có thể làm gì với máy tính?"},
                {"Chủ đề": "MẠNG MÁY TÍNH VÀ INTERNET", "Bài học": "Bài 2. Tìm kiếm thông tin trên website"},
                {"Chủ đề": "TỔ CHỨC LƯU TRỮ, TÌM KIẾM VÀ TRAO ĐỔI THÔNG TIN", "Bài học": "Bài 3. Tìm kiếm thông tin trong giải quyết vấn đề; Bài 4. Cây thư mục"},
                {"Chủ đề": "ĐẠO ĐỨC, PHÁP LUẬT VÀ VĂN HOÁ TRONG MÔI TRƯỜNG SỐ", "Bài học": "Bài 5. Bản quyền nội dung thông tin"},
                {"Chủ đề": "ỨNG DỤNG TIN HỌC", "Bài học": "Bài 6. Định dạng kí tự và bố trí hình ảnh trong văn bản; Bài 7. Thực hành soạn thảo văn bản; Bài 9A: Sử dụng phần mềm đồ họa tạo sản phẩm số; Bài 9B. Thực hành tạo đồ dùng gia đình"}
            ],
            "Học kỳ II": [
                {"Chủ đề": "GIẢI QUYẾT VẤN ĐỀ VỚI SỰ TRỢ GIÚP CỦA MÁY TÍNH", "Bài học": "Bài 10. Cấu trúc tuần tự; Bài 11. Cấu trúc lặp; Bài 12. Thực hành sử dụng lệnh lặp; Bài 13. Cấu trúc rẽ nhánh; Bài 14. Sử dụng biến trong chương trình; Bài 15. Sử dụng biểu thức trong chương trình; Bài 16. Từ kịch bản đến chương trình"}
            ]
        },
        "Công nghệ": {
            "Học kỳ I": [
                {"Chủ đề": "Công nghệ và đời sống", "Bài học": "Bài 1. Vai trò của công nghệ; Bài 2. Nhà sáng chế; Bài 3. Tìm hiểu thiết kế; Bài 4. Thiết kế sản phẩm; Bài 5. Sử dụng điện thoại; Bài 6. Sử dụng tủ lạnh"}
            ],
            "Học kỳ II": [
                {"Chủ đề": "Thủ công kĩ thuật", "Bài học": "Bài 7. Lắp ráp mô hình xe điện chạy bằng pin; Bài 8. Mô hình máy phát điện gió; Bài 9. Mô hình điện mặt trời"}
            ]
        }
    }
}

# --- CẤU TRÚC DỮ LIỆU ĐÃ ĐƯỢC CHUẨN HÓA LẠI ĐỂ TẠO LIST BÀI HỌC ---
CURRICULUM_DB_PROCESSED = {}

# Xử lý dữ liệu thô để tách chuỗi bài học thành list
for grade, subjects in CURRICULUM_DB.items():
    CURRICULUM_DB_PROCESSED[grade] = {}
    for subject, semesters in subjects.items():
        # Xử lý theo từng học kỳ
        CURRICULUM_DB_PROCESSED[grade][subject] = {}
        for semester, content in semesters.items():
            processed_topics = []
            for item in content:
                topic_name = item['Chủ đề']
                raw_lessons_str = item['Bài học']
                # Tách chuỗi dựa trên dấu chấm phẩy
                lessons_list = [l.strip() for l in raw_lessons_str.split(';') if l.strip()]
                
                # Tạo structure mới: mỗi chủ đề chứa một list các bài học con
                processed_topics.append({
                    'Chủ đề': topic_name,
                    'Bài học': lessons_list # Đây giờ là một list các string
                })
            CURRICULUM_DB_PROCESSED[grade][subject][semester] = processed_topics

# --- 5. HỆ THỐNG API MỚI (CHỐNG LỖI 404 VÀ 429) ---
def generate_content_with_rotation(api_key, prompt):
    genai.configure(api_key=api_key)
    
    # 1. LẤY DANH SÁCH MODEL THỰC TẾ TỪ GOOGLE (Tránh lỗi 404 do sai tên)
    try:
        all_models = list(genai.list_models())
    except Exception as e:
        return f"Lỗi kết nối lấy danh sách model: {e}", None

    # Lọc ra các model có thể tạo văn bản
    valid_models = [
        m.name for m in all_models 
        if 'generateContent' in m.supported_generation_methods
    ]
    
    if not valid_models:
        return "Lỗi: API Key đúng nhưng không tìm thấy model nào hỗ trợ tạo văn bản (generateContent).", None

    # 2. SẮP XẾP ƯU TIÊN (Flash > Pro)
    # Chúng ta sẽ tạo một danh sách ưu tiên dựa trên những gì thực tế ĐANG CÓ
    priority_order = []
    
    # Tìm các bản Flash trước
    for m in valid_models:
        if 'flash' in m.lower() and '1.5' in m:
            priority_order.append(m)
            
    # Tìm các bản Pro
    for m in valid_models:
        if 'pro' in m.lower() and '1.5' in m and m not in priority_order:
            priority_order.append(m)
            
    # Các model còn lại (như gemini-pro cũ, gemini-1.0...)
    for m in valid_models:
        if m not in priority_order:
            priority_order.append(m)

    # 3. THỬ LẦN LƯỢT (Cơ chế chống lỗi 429)
    last_error = ""
    
    for model_name in priority_order:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            return response.text, model_name
        except Exception as e:
            error_msg = str(e)
            last_error = error_msg
            # Gặp lỗi (429, 500, v.v.) thì thử model tiếp theo ngay
            time.sleep(1) 
            continue

    return f"Hết model khả dụng. Lỗi cuối cùng: {last_error}", None

# --- 6. HÀM HỖ TRỢ FILE ---
def read_uploaded_file(uploaded_file):
    try:
        if uploaded_file.name.endswith('.xlsx'):
            df = pd.read_excel(uploaded_file)
            return df.to_string()
        elif uploaded_file.name.endswith('.docx'):
            doc = Document(uploaded_file)
            return "\n".join([para.text for para in doc.paragraphs])
        elif uploaded_file.name.endswith('.pdf'):
            if 'pypdf' in globals():
                reader = pypdf.PdfReader(uploaded_file)
                text = ""
                for page in reader.pages: text += page.extract_text()
                return text
        return None
    except Exception:
        return None

def set_font_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

# [YÊU CẦU 5] HÀM TẠO FILE WORD CHO TAB 2 (CÓ MA TRẬN)
def create_word_from_question_list(school_name, subject, exam_list):
    doc = Document()
    set_font_style(doc)
    
    # Header
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(7)
    table.columns[1].width = Cm(9)
    
    cell_1 = table.cell(0, 0)
    p1 = cell_1.paragraphs[0]
    run_s = p1.add_run(f"{school_name.upper()}")
    run_s.bold = True
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cell_2 = table.cell(0, 1)
    p2 = cell_2.paragraphs[0]
    run_e = p2.add_run(f"ĐỀ KIỂM TRA {subject.upper()}\n")
    run_e.bold = True
    run_y = p2.add_run("Năm học: ..........")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    
    # PHẦN 1: MA TRẬN ĐỀ THI (Dựa trên exam_list hiện tại)
    # Lưu ý: Đây là bảng ma trận tóm tắt, bảng chi tiết ở Tab 3 sẽ export riêng
    h1 = doc.add_heading('I. MA TRẬN ĐỀ THI', level=1)
    h1.runs[0].font.name = 'Times New Roman'
    h1.runs[0].font.color.rgb = None
    
    matrix_table = doc.add_table(rows=1, cols=6)
    matrix_table.style = 'Table Grid'
    hdr_cells = matrix_table.rows[0].cells
    headers = ["STT", "Chủ đề / Bài học", "Dạng bài", "Mức độ", "Điểm", "Ghi chú"]
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, q in enumerate(exam_list):
        row_cells = matrix_table.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[1].text = str(q.get('lesson', ''))
        row_cells[2].text = str(q.get('type', ''))
        row_cells[3].text = str(q.get('level', ''))
        row_cells[4].text = str(q.get('points', ''))
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # PHẦN 2: NỘI DUNG ĐỀ THI
    h2 = doc.add_heading('II. NỘI DUNG ĐỀ THI', level=1)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.color.rgb = None
    
    for idx, q in enumerate(exam_list):
        p = doc.add_paragraph()
        run_title = p.add_run(f"Câu {idx + 1} ({q['points']} điểm): ")
        run_title.bold = True
        
        content_lines = q['content'].split('\n')
        for line in content_lines:
            if line.strip():
                if line.startswith("**Câu hỏi:**") or line.startswith("**Đáp án:**"):
                    pass 
                else:
                    doc.add_paragraph(line)
        doc.add_paragraph() 

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# [YÊU CẦU 3 & 4] HÀM TẠO FILE WORD MA TRẬN ĐẶC TẢ (TAB 3)
def create_matrix_document(exam_list, subject_name, grade_name):
    doc = Document()
    
    # Thiết lập khổ giấy ngang cho bảng rộng
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    
    set_font_style(doc)
    
    # Tiêu đề
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"BẢN ĐẶC TẢ ĐỀ KIỂM TRA MÔN {subject_name.upper()} {grade_name.upper()}")
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # Tạo bảng
    # Cấu trúc bảng theo mẫu:
    # Row 1: STT, Chủ đề, Bài, YCCĐ, Dạng câu hỏi, Số điểm, Câu số
    # Row 2 (merged under Dạng CH): TNKQ, TL, Thực hành... (tùy môn)
    # Row 3 (merged under Mức độ): Biết, Hiểu, VD...
    
    # Để đơn giản hoá việc tạo bảng phức tạp bằng python-docx, ta sẽ tạo bảng phẳng 
    # nhưng đầy đủ thông tin các cột như yêu cầu người dùng
    
    table = doc.add_table(rows=2, cols=12)
    table.style = 'Table Grid'
    
    # Header Row 1
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "STT"
    hdr_cells[1].text = "Chủ đề"
    hdr_cells[2].text = "Bài học"
    hdr_cells[3].text = "Yêu cầu cần đạt"
    hdr_cells[4].text = "Dạng câu hỏi & Mức độ nhận thức"
    # Merge cells for "Dạng câu hỏi..." across columns 4 to 10
    hdr_cells[4].merge(hdr_cells[10]) 
    
    hdr_cells[11].text = "Tổng điểm"

    # Header Row 2 (Chi tiết mức độ/dạng)
    row2_cells = table.rows[1].cells
    sub_headers = ["TN-Biết", "TN-Hiểu", "TN-VD", "TL-Biết", "TL-Hiểu", "TL-VD", "Khác"]
    for i, title in enumerate(sub_headers):
        row2_cells[i+4].text = title
        
    # Merge vertical for non-split columns
    for i in [0, 1, 2, 3, 11]:
        hdr_cells[i].merge(row2_cells[i])

    # Fill Data
    # Gom nhóm các câu hỏi theo Chủ đề -> Bài học
    grouped_data = {}
    for idx, q in enumerate(exam_list):
        key = (q['topic'], q['lesson'])
        if key not in grouped_data:
            grouped_data[key] = {'yccd': q.get('yccd', ''), 'questions': []}
        grouped_data[key]['questions'].append(q)

    stt = 1
    for (topic, lesson), data in grouped_data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(stt)
        row_cells[1].text = topic
        row_cells[2].text = lesson
        row_cells[3].text = data['yccd']
        
        # Logic đếm câu hỏi vào các ô mức độ
        # Mapping đơn giản: 
        # TN (Trắc nghiệm, Đúng/Sai, Nối cột, Điền khuyết) -> TN
        # TL (Tự luận) -> TL
        # Mức 1 -> Biết, Mức 2 -> Hiểu, Mức 3 -> VD
        
        counts = {k: [] for k in sub_headers}
        total_points = 0
        
        for q in data['questions']:
            q_idx = exam_list.index(q) + 1
            q_type_code = "TN" if "Tự luận" not in q['type'] and "Thực hành" not in q['type'] else "TL"
            q_level_code = "Biết" if "Mức 1" in q['level'] else ("Hiểu" if "Mức 2" in q['level'] else "VD")
            
            key = f"{q_type_code}-{q_level_code}"
            if key in counts:
                counts[key].append(str(q_idx))
            else:
                 counts["Khác"].append(str(q_idx))
            
            total_points += q['points']
            
        # Fill cells
        for i, key in enumerate(sub_headers):
            if counts[key]:
                row_cells[i+4].text = f"Câu {', '.join(counts[key])}"
        
        row_cells[11].text = str(total_points)
        stt += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_word_file_simple(school_name, exam_name, content):
    doc = Document()
    set_font_style(doc)
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        section.left_margin = Cm(3); section.right_margin = Cm(2)

    table = doc.add_table(rows=1, cols=2); table.autofit = False
    table.columns[0].width = Cm(7); table.columns[1].width = Cm(9)

    cell_1 = table.cell(0, 0); p1 = cell_1.paragraphs[0]
    run_s = p1.add_run(f"{school_name.upper()}"); run_s.bold = True; run_s.font.size = Pt(12)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell_2 = table.cell(0, 1); p2 = cell_2.paragraphs[0]
    run_e = p2.add_run(f"{exam_name.upper()}\n"); run_e.bold = True; run_e.font.size = Pt(12)
    run_y = p2.add_run("Năm học: .........."); run_y.font.size = Pt(13)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    for line in content.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

def extract_periods(lesson_name):
    match = re.search(r'\((\d+)\s*tiết\)', lesson_name, re.IGNORECASE)
    if match:
        return match.group(1)
    return "-"

# --- 7. MAIN APP ---
def main():
    if 'exam_result' not in st.session_state: st.session_state.exam_result = ""
    if "exam_list" not in st.session_state: st.session_state.exam_list = [] 
    if "current_preview" not in st.session_state: st.session_state.current_preview = "" 
    if "temp_question_data" not in st.session_state: st.session_state.temp_question_data = None 
    
    # Biến để kiểm soát việc tự động lấy YCCĐ khi đổi bài học
    if "last_lesson_selected" not in st.session_state: st.session_state.last_lesson_selected = ""
    if "auto_yccd_content" not in st.session_state: st.session_state.auto_yccd_content = "Nắm vững kiến thức cơ bản và vận dụng giải bài tập."

    # --- SIDEBAR CHUNG ---
    with st.sidebar:
        st.header("🔑 CẤU HÌNH HỆ THỐNG")
        st.subheader("HỖ TRỢ RA ĐỀ CẤP TIỂU HỌC")
        api_key = st.text_input("Nhập API Key Google:", type="password")
        
        # NÚT KIỂM TRA API
        if st.button("🔌 Kiểm tra kết nối API"):
            if not api_key:
                st.warning("Vui lòng nhập API Key trước.")
            else:
                try:
                    genai.configure(api_key=api_key)
                    models = list(genai.list_models())
                    st.success(f"✅ Kết nối thành công! (Tìm thấy {len(models)} models)")
                except Exception as e:
                    st.error(f"❌ Kết nối thất bại: {e}")
        
        st.divider()
        st.markdown("**TRƯỜNG PTDTBT TIỂU HỌC GIÀNG CHU PHÌN**")
        st.caption("Hệ thống hỗ trợ chuyên môn")

    if not api_key:
        st.warning("Vui lòng nhập API Key để bắt đầu.")
        return

    # [YÊU CẦU 3] THÊM TIÊU ĐỀ LỚN Ở GIAO DIỆN CHÍNH
    st.markdown('<div class="main-header">HỖ TRỢ RA ĐỀ THI CẤP TIỂU HỌC</div>', unsafe_allow_html=True)

    # --- TABS GIAO DIỆN ---
    tab1, tab2, tab3 = st.tabs(["📁 TẠO ĐỀ TỪ FILE (UPLOAD)", "✍️ SOẠN TỪNG CÂU (CSDL)", "📊 MA TRẬN ĐỀ THI"])

    # ========================== TAB 1: UPLOAD & TẠO ĐỀ ==========================
    with tab1:
        st.header("Tạo đề thi từ file Ma trận có sẵn")
        col1, col2 = st.columns([1, 2])
        with col1:
            st.subheader("1. Chọn Lớp")
            grade_t1 = st.radio("Khối lớp:", list(SUBJECTS_DB.keys()), key="t1_grade")
        with col2:
            st.subheader("2. Chọn Môn")
            subjects_t1 = SUBJECTS_DB[grade_t1]
            sub_name_t1 = st.selectbox("Môn học:", [s[0] for s in subjects_t1], key="t1_sub")
            icon_t1 = next(i for n, i in subjects_t1 if n == sub_name_t1)
            st.markdown(f"<div class='subject-card'><h3>{icon_t1} {sub_name_t1}</h3></div>", unsafe_allow_html=True)
            exam_term_t1 = st.selectbox("Kỳ thi:", 
                ["ĐỀ KIỂM TRA ĐỊNH KÌ GIỮA HỌC KÌ I", "ĐỀ KIỂM TRA ĐỊNH KÌ CUỐI HỌC KÌ I",
                "ĐỀ KIỂM TRA ĐỊNH KÌ GIỮA HỌC KÌ II", "ĐỀ KIỂM TRA ĐỊNH KÌ CUỐI HỌC KÌ II"], key="t1_term")
            school_name_t1 = st.text_input("Tên trường:", value="TRƯỜNG PTDTBT TIỂU HỌC GIÀNG CHU PHÌN", key="t1_school")

        st.subheader("3. Upload Ma trận")
        # [YÊU CẦU 1: Bỏ dòng hướng dẫn có bóng đèn]
        uploaded = st.file_uploader("Chọn file (.xlsx, .docx, .pdf)", type=['xlsx', 'docx', 'pdf'], key="t1_up")

        if uploaded and st.button("🚀 TẠO ĐỀ THI NGAY", type="primary", key="t1_btn"):
            content = read_uploaded_file(uploaded)
            if content:
                with st.spinner("Đang phân tích ma trận và tạo đề từ nguồn GDPT 2018..."):
                    # [YÊU CẦU 1 SỬA LẠI: PHÂN TÍCH FILE ĐỂ TÌM BỘ SÁCH VÀ TÍNH ĐIỂM CHÍNH XÁC]
                    prompt = f"""
                    Bạn là chuyên gia giáo dục tiểu học Việt Nam.
                    Nhiệm vụ: Soạn đề thi môn {sub_name_t1} lớp {grade_t1} dựa CHÍNH XÁC vào file ma trận tải lên.

                    PHÂN TÍCH DỮ LIỆU ĐẦU VÀO (QUAN TRỌNG):
                    1. Xác định bộ sách giáo khoa: Đọc file để tìm từ khóa (Chân trời sáng tạo, Kết nối tri thức, Cánh diều, Cùng khám phá...). Nếu không thấy, hãy dùng bộ sách phổ biến nhất cho {sub_name_t1} lớp {grade_t1}.
                    2. Phân tích điểm số logic:
                       - Nếu ma trận ghi "Tổng điểm" cho một hàng có nhiều câu hỏi (ví dụ: Số câu: 2, Tổng điểm: 1.0), thì điểm mỗi câu = Tổng điểm / Số câu = 0.5 điểm. 
                       - TUYỆT ĐỐI KHÔNG GÁN tổng điểm (ví dụ 25 điểm) cho 1 câu hỏi trắc nghiệm đơn lẻ. Điểm mỗi câu trắc nghiệm thường là 0.5 hoặc 1.0.
                    
                    QUY TRÌNH TẠO ĐỀ:
                    1. Tạo đúng số lượng câu hỏi theo ma trận.
                    2. Nội dung câu hỏi phải cụ thể, rõ ràng, có dữ kiện đầy đủ (Không được viết chung chung kiểu "Hãy làm bài tập này").
                    3. Đối với môn Tin học/Công nghệ: Câu hỏi phải bám sát phần mềm/công cụ dạy trong chương trình 2018.

                    ĐỊNH DẠNG ĐẦU RA BẮT BUỘC:
                    **Câu [Số thứ tự]** ([Điểm đã chia nhỏ] đ) - [Mức độ]: [Nội dung câu hỏi đầy đủ]
                    A. [Nội dung lựa chọn A]
                    B. [Nội dung lựa chọn B]
                    C. [Nội dung lựa chọn C]
                    D. [Nội dung lựa chọn D]
                    (Xuống dòng) Đáp án: [Đáp án đúng]

                    DỮ LIỆU TỪ FILE MA TRẬN:
                    {content}
                    """
                    result_text, used_model = generate_content_with_rotation(api_key, prompt)
                    if used_model:
                        st.session_state.exam_result = result_text
                        st.success(f"Đã tạo xong bằng model: {used_model}")
                    else:
                        st.error(result_text)

        if st.session_state.exam_result:
            st.markdown("---")
            edited_text = st.text_area("Sửa nội dung:", value=st.session_state.exam_result, height=500, key="t1_edit")
            st.session_state.exam_result = edited_text 
            docx = create_word_file_simple(school_name_t1, exam_term_t1, edited_text)
            st.download_button("📥 TẢI VỀ FILE WORD (.docx)", docx, file_name=f"De_{sub_name_t1}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")

    # ========================== TAB 2: SOẠN TỪNG CÂU ==========================
    with tab2:
        st.header("Soạn thảo từng câu hỏi theo CSDL")
        col1, col2 = st.columns(2)
        with col1:
            selected_grade = st.selectbox("Chọn Khối Lớp:", list(SUBJECTS_DB.keys()), key="t2_grade")
        with col2:
            subjects_list = [f"{s[1]} {s[0]}" for s in SUBJECTS_DB[selected_grade]]
            selected_subject_full = st.selectbox("Chọn Môn Học:", subjects_list, key="t2_sub")
            selected_subject = selected_subject_full.split(" ", 1)[1]

        # Lấy dữ liệu đã được xử lý từ CURRICULUM_DB_PROCESSED
        raw_data = CURRICULUM_DB_PROCESSED.get(selected_grade, {}).get(selected_subject, {})

        if not raw_data:
            st.warning("⚠️ Dữ liệu môn này đang cập nhật.")
        else:
            st.markdown("---")
            st.subheader("🛠️ Soạn thảo câu hỏi")
            
            col_a, col_b = st.columns(2)
            with col_a:
                # Xử lý Dropdown Học kỳ
                all_terms = list(raw_data.keys())
                selected_term = st.selectbox("Chọn Học kỳ:", all_terms, key="t2_term")
                
                # Lấy danh sách các bài học (bao gồm chủ đề) thuộc học kỳ đã chọn
                lessons_in_term = raw_data[selected_term]

                # Gom danh sách chủ đề duy nhất
                unique_topics = sorted(list(set([l['Chủ đề'] for l in lessons_in_term])))
                selected_topic = st.selectbox("Chọn Chủ đề:", unique_topics, key="t2_topic")

            with col_b:
                # Lọc bài học THEO CHỦ ĐỀ ĐÃ CHỌN
                filtered_lessons = [l for l in lessons_in_term if l['Chủ đề'] == selected_topic]
                
                # Tạo list tất cả bài học nhỏ
                all_lessons_in_topic = []
                for item in filtered_lessons:
                    # item['Bài học'] bây giờ là một LIST các chuỗi bài học nhỏ
                    all_lessons_in_topic.extend(item['Bài học'])
                
                # Dropdown chọn bài học
                selected_lesson_name = st.selectbox("Chọn Bài học:", all_lessons_in_topic, key="t2_lesson")
                
                # [YÊU CẦU 2: SỬA LỖI HÌNH 22222: YCCĐ LẤY ĐÚNG THEO CHUẨN]
                if st.session_state.last_lesson_selected != selected_lesson_name:
                    with st.spinner("Đang tra cứu YCCĐ chuẩn GDPT 2018 (Chế độ chuyên gia)..."):
                        # Prompt được tối ưu để đóng vai ChatGPT chuyên gia
                        yccd_prompt = f"""
                        AI đang chạy
                        Nhiệm vụ: Trích xuất chính xác Yêu cầu cần đạt (YCCĐ) cho bài học sau:
                        - Bài học: '{selected_lesson_name}'
                        - Chủ đề: '{selected_topic}'
                        - Môn: {selected_subject}
                        - Lớp: {selected_grade}
                        
                        Yêu cầu:
                        1. Chỉ đưa ra nội dung cốt lõi, ngắn gọn, súc tích.
                        2. Phải chính xác với văn bản quy định của Bộ GD&ĐT.
                        3. Không thêm lời dẫn.
                        """
                        ai_yccd, _ = generate_content_with_rotation(api_key, yccd_prompt)
                        if ai_yccd:
                            st.session_state.auto_yccd_content = ai_yccd
                        st.session_state.last_lesson_selected = selected_lesson_name
                
                # Input YCCĐ (Hiển thị giá trị từ session state)
                yccd_input = st.text_area("Yêu cầu cần đạt (AI tự động lấy):", value=st.session_state.auto_yccd_content, height=68, key="t2_yccd_input")
                
                # Lưu thông tin bài học hiện tại để dùng
                current_lesson_data = {
                    "Chủ đề": selected_topic,
                    "Bài học": selected_lesson_name,
                    "YCCĐ": yccd_input
                }

            col_x, col_y, col_z = st.columns(3)
            with col_x:
                # [YÊU CẦU 2: Sửa tên dạng câu hỏi cho đúng thực tế]
                question_types = [
                    "Trắc nghiệm (4 lựa chọn)", 
                    "Đúng/Sai", 
                    "Ghép nối (Nối cột)", 
                    "Điền khuyết (Hoàn thành câu)", 
                    "Tự luận"
                ]
                if selected_subject == "Tin học":
                    question_types.append("Thực hành trên máy tính")
                    
                q_type = st.selectbox("Dạng câu hỏi:", question_types, key="t2_type")
            with col_y:
                level = st.selectbox("Mức độ:", ["Mức 1: Biết", "Mức 2: Hiểu", "Mức 3: Vận dụng"], key="t2_lv")
            with col_z:
                points = st.number_input("Điểm số:", min_value=0.25, max_value=10.0, step=0.25, value=1.0, key="t2_pt")

            # HÀM TẠO CÂU HỎI
            def generate_question():
                with st.spinner("AI đang viết..."):
                    random_seed = random.randint(1, 100000)
                    # [SỬA LỖI HÌNH 3333 VÀ 4444: TỐI ƯU HÓA PROMPT TẠO CÂU HỎI]
                    prompt_q = f"""
                    Đóng vai chuyên gia giáo dục Tiểu học. Soạn **1 CÂU HỎI KIỂM TRA** môn {selected_subject} Lớp {selected_grade}.
                    - Chủ đề: {current_lesson_data['Chủ đề']}
                    - Bài học cụ thể: {current_lesson_data['Bài học']}
                    - YCCĐ: {current_lesson_data['YCCĐ']}
                    - Dạng: {q_type} - Mức độ: {level} - Điểm: {points}
                    - Seed ngẫu nhiên: {random_seed}
                    
                    YÊU CẦU ĐỊNH DẠNG NGHIÊM NGẶT (SỬA LỖI HIỂN THỊ):
                    1. VỚI DẠNG "Trắc nghiệm (4 lựa chọn)":
                    - Phải hiển thị 4 đáp án A. B. C. D. riêng biệt xuống dòng.
                    - Chỉ ra đáp án đúng ở cuối.
                    2. VỚI DẠNG "Ghép nối (Nối cột)":
                    - Phải liệt kê nội dung Cột A (1, 2,...) và Cột B (a, b,...) rõ ràng.
                    - Phần đáp án mô phỏng kết quả nối (ví dụ: 1-b, 2-a).
                    3. VỚI DẠNG "Điền khuyết" hoặc "Tự luận":
                    - Câu hỏi phải chừa chỗ trống bằng dấu ".........." để học sinh điền.
                    - Hiển thị đáp án gợi ý ở cuối.

                    OUTPUT CHỈ GHI NỘI DUNG, KHÔNG CẦN LỜI DẪN:
                    [Nội dung câu hỏi và các lựa chọn]
                    
                    Đáp án: ...
                    """
                    preview_content, _ = generate_content_with_rotation(api_key, prompt_q)
                    st.session_state.current_preview = preview_content
                    st.session_state.temp_question_data = {
                        "topic": selected_topic, "lesson": selected_lesson_name,
                        "type": q_type, "level": level, "points": points, "content": preview_content,
                        "yccd": yccd_input, "periods": extract_periods(selected_lesson_name)
                    }

            if st.button("✨ Tạo câu hỏi (Xem trước)", type="primary", key="t2_preview"):
                generate_question()

            if st.session_state.current_preview:
                st.markdown(f"<div class='question-box'>{st.session_state.current_preview}</div>", unsafe_allow_html=True)
                
                col_btn1, col_btn2 = st.columns([1, 1])
                with col_btn1:
                    if st.button("✅ Thêm vào đề thi", key="t2_add"):
                        st.session_state.exam_list.append(st.session_state.temp_question_data)
                        st.session_state.current_preview = ""
                        st.success("Đã thêm vào danh sách!")
                        st.rerun()
                with col_btn2:
                    if st.button("🔄 Tạo câu hỏi khác", key="t2_regen"):
                        generate_question()
                        st.rerun()

            # --- DANH SÁCH & THỐNG KÊ ---
            if len(st.session_state.exam_list) > 0:
                st.markdown("---")
                
                st.subheader(f"📊 Bảng thống kê chi tiết ({len(st.session_state.exam_list)} câu)")
                
                # Chuẩn bị dữ liệu cho bảng
                stats_data = []
                for i, q in enumerate(st.session_state.exam_list):
                    stats_data.append({
                        "Thứ tự câu": f"Câu {i+1}",
                        "Tên bài": q['lesson'],
                        "Số tiết": q.get('periods', '-'),
                        "Các mức": q['level'],
                        "Dạng câu hỏi": q['type'],
                        "Điểm": q['points']
                    })
                
                df_stats = pd.DataFrame(stats_data)
                st.dataframe(df_stats, use_container_width=True)

                st.markdown("#### 📝 Chỉnh sửa chi tiết đề thi")
                # Duyệt qua từng câu để hiển thị ô nhập liệu
                for i, item in enumerate(st.session_state.exam_list):
                    with st.expander(f"Câu {i+1} ({item['points']} điểm) - {item['type']}"):
                        new_content = st.text_area(
                            f"Nội dung câu {i+1}:", 
                            value=item['content'], 
                            height=150, 
                            key=f"edit_q_{i}"
                        )
                        st.session_state.exam_list[i]['content'] = new_content
                        
                        if st.button("🗑️ Xóa câu này", key=f"del_q_{i}"):
                            st.session_state.exam_list.pop(i)
                            st.rerun()

                col_act1, col_act2 = st.columns(2)
                with col_act2:
                     if st.button("❌ Xóa toàn bộ đề", key="t2_clear"):
                        st.session_state.exam_list = []
                        st.rerun()

                docx_file = create_word_from_question_list("TRƯỜNG PTDTBT TIỂU HỌC GIÀNG CHU PHÌN", selected_subject, st.session_state.exam_list)
                st.download_button(
                    label="📥 TẢI ĐỀ THI (WORD)", 
                    data=docx_file,
                    file_name=f"De_thi_{selected_subject}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )
    
    # ========================== TAB 3: MA TRẬN ĐỀ THI ==========================
    with tab3:
        st.header("📊 BẢNG MA TRẬN ĐỀ THI (BẢN ĐẶC TẢ)")
        st.info("Chỉnh sửa trực tiếp trên bảng và tải về file Word theo mẫu.")
        
        if len(st.session_state.exam_list) == 0:
            st.info("⚠️ Vui lòng soạn câu hỏi ở Tab 2 trước.")
        else:
            # Tạo DataFrame cho Ma trận từ danh sách câu hỏi
            matrix_data = []
            for i, q in enumerate(st.session_state.exam_list):
                matrix_data.append({
                    "STT": i + 1,
                    "Chủ đề": q['topic'],
                    "Bài học": q['lesson'],
                    "Yêu cầu cần đạt": q.get('yccd', ''),
                    "Dạng câu hỏi": q['type'],
                    "Mức độ": q['level'],
                    "Số điểm": q['points'],
                    "Ghi chú": ""
                })
            
            df_matrix = pd.DataFrame(matrix_data)
            
            # Hiển thị bảng Data Editor để chỉnh sửa trực tiếp
            edited_df = st.data_editor(
                df_matrix,
                num_rows="dynamic",
                use_container_width=True,
                key="matrix_editor"
            )
            
            # Cập nhật ngược lại session_state nếu có thay đổi quan trọng
            if st.button("💾 Cập nhật thay đổi từ Ma trận vào Hệ thống"):
                for index, row in edited_df.iterrows():
                    if index < len(st.session_state.exam_list):
                        st.session_state.exam_list[index]['topic'] = row['Chủ đề']
                        st.session_state.exam_list[index]['lesson'] = row['Bài học']
                        st.session_state.exam_list[index]['type'] = row['Dạng câu hỏi']
                        st.session_state.exam_list[index]['level'] = row['Mức độ']
                        st.session_state.exam_list[index]['points'] = row['Số điểm']
                        st.session_state.exam_list[index]['yccd'] = row['Yêu cầu cần đạt']
                st.success("Đã cập nhật dữ liệu thành công!")
                st.rerun()

            # [YÊU CẦU 3 & 4] TẢI WORD MA TRẬN THEO MẪU
            matrix_docx = create_matrix_document(st.session_state.exam_list, selected_subject, selected_grade)
            st.download_button(
                label="📥 TẢI BẢN ĐẶC TẢ ĐỀ THI (WORD)",
                data=matrix_docx,
                file_name=f"Ban_dac_ta_{selected_subject}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )

    # --- FOOTER ---
    st.markdown("""
    <div class="footer">
        <p style="margin: 0; font-weight: bold; color: #2c3e50;">🏫 TRƯỜNG PTDTBT TIỂU HỌC GIÀNG CHU PHÌN</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
